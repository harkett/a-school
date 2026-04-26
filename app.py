import io
import streamlit as st
from pathlib import Path
from datetime import date

from src.prompts import build_prompt
from src.generator import generate, transcribe_audio, transcribe_image
from src.formats import to_markdown, save
from src.auth import (generate_magic_token, verify_magic_token, peek_magic_token,
                      send_magic_link, get_current_user, notify_admin_connexion,
                      create_session, get_session, delete_session)
from streamlit_cookies_controller import CookieController
from streamlit_mic_recorder import mic_recorder

st.set_page_config(page_title="A-SCHOOL — Générateur pédagogique", page_icon="📚", layout="wide")

_cookies = CookieController()


def _get_webmail_url(email: str) -> str | None:
    domain = email.split("@")[-1].lower()
    if domain == "gmail.com":
        return "https://mail.google.com"
    if domain in ("outlook.fr", "outlook.com", "hotmail.com", "hotmail.fr", "live.fr", "live.com"):
        return "https://outlook.live.com/mail/inbox"
    if domain in ("yahoo.fr", "yahoo.com"):
        return "https://mail.yahoo.com"
    if domain == "laposte.net":
        return "https://www.laposte.net/accueil"
    if domain == "orange.fr":
        return "https://messagerie.orange.fr"
    return None

# ── Restauration session depuis cookie ───────────────────────────────────────
# CookieController est asynchrone : les cookies ne sont disponibles qu'au 2e cycle.
# On force un rerun unique au démarrage pour les charger avant d'afficher le login.
if not st.session_state.get("user_email"):
    try:
        _session_token = _cookies.get("aschool_session")
    except Exception:
        _session_token = None
    if _session_token:
        _session_data = get_session(_session_token)
        if _session_data:
            st.session_state["user_email"] = _session_data["email"]
            st.session_state["matiere"] = _session_data["matiere"]
            st.session_state["session_token"] = _session_token
    elif not st.session_state.get("_cookie_init_done"):
        st.session_state["_cookie_init_done"] = True
        st.rerun()

# ── Logout via URL ────────────────────────────────────────────────────────────
if "logout" in st.query_params:
    _token = st.session_state.pop("session_token", None)
    if _token:
        delete_session(_token)
        _cookies.remove("aschool_session")
    st.session_state.pop("user_email", None)
    st.session_state.pop("user_name", None)
    st.query_params.clear()
    st.rerun()

# ── Vérification token Magic link dans l'URL ─────────────────────────────────
params = st.query_params
if "token" in params and not st.session_state.get("user_email"):
    result = verify_magic_token(params["token"])
    if result:
        st.session_state["user_email"] = result["email"]
        st.session_state["matiere"] = result["matiere"]
        _tok = create_session(result["email"], result["matiere"])
        st.session_state["session_token"] = _tok
        _cookies.set("aschool_session", _tok, max_age=30 * 24 * 3600)
        st.query_params.clear()
        try:
            notify_admin_connexion(result["email"], "A-SCHOOL : Générateur d'activités pédagogiques")
        except Exception:
            pass
        st.rerun()
    else:
        st.error("Lien invalide ou expiré. Demandez un nouveau lien.")
        st.stop()
    st.rerun()

# ── Page de connexion ─────────────────────────────────────────────────────────
user = get_current_user()

if user and user["method"] == "google" and not st.session_state.get("google_notified"):
    try:
        notify_admin_connexion(user["email"], "google")
        st.session_state["google_notified"] = True
    except Exception:
        pass

if not user:
    st.markdown("""
    <style>
        .block-container { padding-top: 5rem; max-width: 480px; }
        #MainMenu, header, [data-testid="stToolbar"],
        [data-testid="stSidebar"], [data-testid="collapsedControl"] { display: none !important; }
        .aschool-navbar {
            position: fixed; top: 0; left: 0; right: 0; z-index: 1000;
            background: #1F6EEB;
            display: flex; align-items: center; justify-content: space-between;
            padding: 0 2rem; height: 56px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.2);
        }
        .aschool-navbar-logo { font-size:1.1rem; font-weight:700; color:white; letter-spacing:-0.3px; }
        .aschool-navbar-logo span { color:#A63045; }
        .aschool-navbar-btn {
            color: white !important; text-decoration: none !important;
            border: 1px solid rgba(255,255,255,0.35); border-radius: 6px;
            padding: 0.3rem 1rem; font-size: 0.875rem; font-weight: 500;
        }
        .aschool-navbar-btn:hover { background: rgba(255,255,255,0.1); }
    </style>
    <nav class="aschool-navbar">
        <div class="aschool-navbar-logo"><span>A</span>-SCHOOL</div>
        <a href="#" class="aschool-navbar-btn" title="Entrez votre email pour recevoir un lien de connexion">Connexion</a>
    </nav>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style="background:linear-gradient(135deg,#1558C0,#1F6EEB);
                border-radius:12px; padding:2rem; text-align:center; margin-bottom:2rem;
                max-width:420px; margin-left:auto; margin-right:auto;">
        <h1 style="color:white; margin:0; font-size:2rem;"><span style="color:#A63045;">A</span>-SCHOOL</h1>
        <p style="color:rgba(255,255,255,0.85); margin:0.5rem 0 0 0;">
            Générateur d'activités pédagogiques
        </p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("Entrez votre email pour recevoir un lien de connexion :")

    with st.form("magic_link_form"):
        email_input = st.text_input("Votre adresse email professionnelle",
                                    placeholder="prenom.nom@ac-bordeaux.fr")
        matiere_input = st.selectbox("Vous enseignez quelle matière ?",
                                     ["Français", "Histoire-Géographie"])
        envoyer = st.form_submit_button("Recevoir un lien de connexion",
                                        use_container_width=True)

    if envoyer:
        if not email_input or "@" not in email_input:
            st.error("Entrez une adresse email valide.")
        else:
            with st.spinner("Envoi en cours..."):
                try:
                    token = generate_magic_token(email_input, matiere_input)
                    send_magic_link(email_input, token)
                    webmail = _get_webmail_url(email_input)
                    email_display = f'<a href="{webmail}" target="_blank">{email_input}</a>' if webmail else email_input
                    st.markdown(
                        f'<div style="background:#f0fdf4;border:1px solid #bbf7d0;border-radius:8px;padding:0.75rem 1rem;color:#166534;">'
                        f'Lien envoyé à {email_display} — vérifiez votre boîte mail (valable 15 min).</div>',
                        unsafe_allow_html=True
                    )
                except Exception as e:
                    st.error(f"Erreur d'envoi : {e}")
    st.stop()

# ── Styles ───────────────────────────────────────────────────────────────────
st.markdown("""
<style>
    .block-container { padding-top: 4.5rem; padding-bottom: 4.5rem; max-width: 960px; }
    #MainMenu, header, [data-testid="stToolbar"] { display: none !important; }
    [data-testid="stSidebar"] { display: none !important; }
    [data-testid="collapsedControl"] { display: none !important; }

    /* Navbar */
    .aschool-navbar {
        position: fixed; top: 0; left: 0; right: 0; z-index: 1000;
        background: #1F6EEB;
        display: flex; align-items: center; justify-content: space-between;
        padding: 0 2rem; height: 56px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.2);
    }
    .aschool-navbar-left { display: flex; align-items: center; gap: 0.75rem; }
    .aschool-navbar-logo { font-size: 1.1rem; font-weight: 700; color: white; letter-spacing: -0.3px; }
    .aschool-navbar-logo span { color: #A63045; }
    .aschool-navbar-slogan { color: rgba(255,255,255,0.65); font-size: 0.875rem; }
    .aschool-navbar-right { display: flex; align-items: center; gap: 1.5rem; font-size: 0.875rem; }
    .aschool-navbar-email { color: rgba(255,255,255,0.6); }
    .aschool-navbar-matiere { color: #A63045; font-weight: 600; }
    .aschool-navbar-logout {
        color: white !important; text-decoration: none !important;
        border: 1px solid rgba(255,255,255,0.35); border-radius: 6px;
        padding: 0.3rem 1rem; font-weight: 500;
        display: inline-flex; align-items: center; gap: 5px;
        transition: background 0.15s;
    }
    .aschool-navbar-logout:hover { background: rgba(255,255,255,0.1); }

    /* Section title */
    .aschool-section-title {
        border-left: 3px solid #A63045;
        padding-left: 8px;
        font-size: 0.75rem; font-weight: 700;
        text-transform: uppercase; letter-spacing: 0.6px;
        color: #64748b;
        margin-bottom: 1rem;
    }

    /* Séparateurs */
    hr { border-color: #e2e8f0 !important; margin: 1.25rem 0 !important; }

    /* Résultat */
    .aschool-result {
        background: #f8faff;
        border: 1px solid #e2e8f0;
        border-left: 4px solid #A63045;
        border-radius: 8px;
        padding: 1.5rem;
        margin-top: 1rem;
    }

    /* Bouton Générer */
    [data-testid="stButton"] > button[kind="primary"] {
        background: #1F6EEB !important;
        border: none !important; border-radius: 6px !important;
        font-size: 1rem !important; font-weight: 600 !important;
        letter-spacing: 0.3px !important; padding: 0.6rem 1.5rem !important;
        box-shadow: 0 2px 8px rgba(31,110,235,0.3) !important;
        transition: all 0.15s ease !important;
    }
    [data-testid="stButton"] > button[kind="primary"]:hover {
        background: #1558C0 !important;
    }

    /* Boutons téléchargement */
    [data-testid="stDownloadButton"] > button {
        border-radius: 6px !important;
        border-color: #A63045 !important;
        color: #A63045 !important;
        font-weight: 500 !important;
    }

    /* File uploader */
    [data-testid="stFileUploader"] label { display: none; }

    /* Selectbox */
    [data-testid="stSelectbox"] input { pointer-events: none !important; caret-color: transparent !important; }
    [data-testid="stSelectbox"] > div { cursor: pointer !important; }
    [data-testid="stSelectbox"] > div > div { cursor: pointer !important; border-radius: 6px !important; border-color: #cbd5e1 !important; }

    /* Focus bordeaux */
    [data-testid="stSelectbox"] > div > div:focus-within { border-color: #A63045 !important; }

    /* Captions */
    [data-testid="stCaptionContainer"] { color: #94a3b8 !important; font-size: 0.8rem !important; }

    /* Zone micro */
    [data-testid="stColumn"]:has(audio) [data-testid="stVerticalBlock"] {
        background-color: #f0f4f8; border-radius: 8px;
        padding: 0.75rem 1rem; min-height: 88px;
    }
    [data-testid="stColumn"]:has(audio) button {
        background-color: white !important; color: #1e293b !important;
        border: 1px solid #cbd5e1 !important; border-radius: 6px !important;
        padding: 0.25rem 0.75rem !important; font-size: 0.875rem !important;
        font-weight: 400 !important; box-shadow: none !important;
    }

    /* Footer */
    .aschool-footer {
        position: fixed; bottom: 0; left: 0; right: 0;
        background: white; border-top: 1px solid #e2e8f0;
        padding: 0.4rem 2rem; z-index: 999;
    }
</style>
""", unsafe_allow_html=True)

# ── Navbar ────────────────────────────────────────────────────────────────────
matiere_nav = st.session_state.get("matiere", "Français")
st.markdown(f"""
<nav class="aschool-navbar">
    <div class="aschool-navbar-left">
        <div class="aschool-navbar-logo"><span>A</span>-SCHOOL</div>
        <span class="aschool-navbar-slogan">| Générateur d'activités pédagogiques</span>
    </div>
    <div class="aschool-navbar-right">
        <span style="color:rgba(255,255,255,0.3);">|</span>
        <span class="aschool-navbar-matiere">{matiere_nav}</span>
        <span style="color:rgba(255,255,255,0.3);">|</span>
        <span class="aschool-navbar-email">{user['email']}</span>
        <a href="?logout=1" class="aschool-navbar-logout"
           title="Fermer votre session et revenir à la page de connexion">
            <svg xmlns="http://www.w3.org/2000/svg" width="13" height="13" viewBox="0 0 24 24"
                 fill="none" stroke="currentColor" stroke-width="2">
                <path d="M9 21H5a2 2 0 0 1-2-2V5a2 2 0 0 1 2-2h4"/>
                <polyline points="16 17 21 12 16 7"/>
                <line x1="21" y1="12" x2="9" y2="12"/>
            </svg>
            Se déconnecter
        </a>
    </div>
</nav>
""", unsafe_allow_html=True)

# ── Footer ────────────────────────────────────────────────────────────────────
st.markdown("""
<footer class="aschool-footer">
    <p style="text-align:center; font-size:0.72rem; font-style:italic; color:#94a3b8; margin:0 0 0.15rem 0;">
        "Plus le monde s'ouvre, plus nous avons besoin de proximité..."
    </p>
    <div style="display:flex; justify-content:space-between; font-size:0.7rem; color:#94a3b8;">
        <span><span style="color:#A63045; font-weight:600;">A</span>-SCHOOL
              &nbsp;·&nbsp;<a href="#" style="color:#94a3b8; text-decoration:none;">Mentions légales</a></span>
        <span>v3.0-dev &nbsp;·&nbsp; harketti@afia.fr</span>
    </div>
</footer>
""", unsafe_allow_html=True)

# ── Configuration des activités ───────────────────────────────────────────────
ACTIVITES_PAR_MATIERE = {}

ACTIVITES_PAR_MATIERE["Français"] = {
    "Questions de compréhension": {
        "key": "comprehension",
        "sous_types": ["simples (repérage)", "inférence", "interprétation", "personnages", "décor / contexte", "émotions / intentions", "mélange"],
        "params": ["nb", "sous_type"],
    },
    "Pistes de lecture": {
        "key": "pistes",
        "sous_types": ["thématique", "narrateur", "réalisme", "point de vue", "registre"],
        "params": ["nb", "sous_type"],
    },
    "Résumés": {
        "key": "resume",
        "sous_types": ["court (5 lignes)", "structuré (début/milieu/fin)", "pour l'oral"],
        "params": ["sous_type"],
    },
    "Analyse de texte": {
        "key": "analyse",
        "sous_types": ["thème principal", "champs lexicaux", "procédés d'écriture", "passage difficile"],
        "params": ["sous_type"],
    },
    "Exercices de réécriture": {
        "key": "reecriture",
        "sous_types": [
            "style direct vers style indirect",
            "passé simple vers présent",
            "présent vers passé simple",
            "présent vers conditionnel présent",
            "1ère personne du singulier vers 1ère personne du pluriel",
            "3ème personne du singulier vers 1ère personne du pluriel",
            "changement de point de vue",
            "simplifier le vocabulaire",
            "enrichir le texte",
        ],
        "params": ["sous_type"],
    },
    "Étude de vocabulaire": {
        "key": "vocabulaire",
        "sous_types": ["mots difficiles + définitions", "synonymes / antonymes", "exercices à trous", "reformulation"],
        "params": ["sous_type"],
    },
    "Production d'écrit": {
        "key": "production_ecrit",
        "sous_types": ["paragraphe argumenté", "continuer le texte", "décrire un personnage", "imaginer la suite d'une scène", "texte poétique"],
        "params": ["sous_type"],
    },
    "Questions pour l'oral": {
        "key": "oral",
        "sous_types": ["débat", "exposé", "échange en classe"],
        "params": ["nb", "sous_type"],
    },
    "Fiche pédagogique": {
        "key": "fiche_pedagogique",
        "sous_types": [],
        "params": [],
    },
    "Exercices de grammaire": {
        "key": "grammaire",
        "sous_types": ["temps verbaux", "types de phrases", "transformer des phrases", "accords"],
        "params": ["sous_type"],
    },
    "Recherche de séquences": {
        "key": "recherche_sequences",
        "sous_types": [],
        "params": [],
    },
    "Séquence détaillée": {
        "key": "sequence_detaillee",
        "sous_types": [],
        "params": [],
    },
    "Questionnaire sur un roman": {
        "key": "questionnaire_roman",
        "sous_types": [],
        "params": [],
    },
    "Évaluation de grammaire": {
        "key": "evaluation_grammaire",
        "sous_types": [],
        "params": [],
    },
    "Évaluation d'orthographe": {
        "key": "evaluation_orthographe",
        "sous_types": [],
        "params": [],
    },
}

ACTIVITES_PAR_MATIERE["Histoire-Géographie"] = {
    "Questions sur un document": {
        "key": "hg_comprehension",
        "sous_types": ["identification", "contexte", "analyse", "mise en relation", "mélange"],
        "params": ["nb", "sous_type"],
    },
    "Analyse de source": {
        "key": "hg_analyse_source",
        "sous_types": [],
        "params": [],
    },
    "Questions de cours": {
        "key": "hg_questions_cours",
        "sous_types": ["connaissances", "définitions", "explication", "mélange"],
        "params": ["nb", "sous_type"],
    },
    "Frise chronologique": {
        "key": "hg_frise",
        "sous_types": [],
        "params": [],
    },
    "Paragraphe argumenté": {
        "key": "hg_paragraphe",
        "sous_types": ["réponse organisée", "SEUL", "bilan de séquence"],
        "params": ["sous_type"],
    },
    "Fiche de révision": {
        "key": "hg_fiche_revision",
        "sous_types": [],
        "params": [],
    },
    "Composition / Dissertation": {
        "key": "hg_composition",
        "sous_types": ["introduction seule", "plan détaillé", "développement complet", "plan avec transitions"],
        "params": ["sous_type"],
    },
    "Lecture de carte / Croquis": {
        "key": "hg_carte",
        "sous_types": ["décrire et expliquer une carte", "questions sur un croquis", "légende à compléter"],
        "params": ["sous_type"],
    },
    "Étude d'un document iconographique": {
        "key": "hg_iconographie",
        "sous_types": ["affiche de propagande", "dessin de presse", "photographie historique", "œuvre d'art"],
        "params": ["sous_type"],
    },
    "Exercice de repères": {
        "key": "hg_reperes",
        "sous_types": ["QCM de repères", "définir des notions clés", "placer des événements sur une frise", "situer des lieux"],
        "params": ["nb", "sous_type"],
    },
    "Mise en relation de documents": {
        "key": "hg_mise_en_relation",
        "sous_types": ["confronter deux sources", "dégager complémentarité / contradiction", "synthèse de documents"],
        "params": ["sous_type"],
    },
    "Préparation à l'oral": {
        "key": "hg_oral",
        "sous_types": ["exposé", "débat", "Grand Oral Terminale", "échange en classe"],
        "params": ["nb", "sous_type"],
    },
}


def to_docx(texte: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)
    for ligne in texte.split("\n"):
        ligne = ligne.strip()
        if ligne.startswith("# "):
            doc.add_heading(ligne[2:], level=1)
        elif ligne.startswith("## "):
            doc.add_heading(ligne[3:], level=2)
        elif ligne.startswith("### "):
            doc.add_heading(ligne[4:], level=3)
        elif ligne:
            doc.add_paragraph(ligne)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Section Texte source ──────────────────────────────────────────────────────
st.markdown('<div class="aschool-section-title">Texte source</div>', unsafe_allow_html=True)

c_upload, c_micro = st.columns([1, 1])
with c_upload:
    fichier = st.file_uploader(
        "Importer",
        type=["txt", "jpg", "jpeg", "png"],
        label_visibility="collapsed",
        help="Fichier texte (.txt) ou image scannée (JPG/PNG) — le texte sera extrait automatiquement",
    )
    st.caption("Texte .txt ou scan JPG/PNG")
with c_micro:
    audio = mic_recorder(
        start_prompt="Dicter",
        stop_prompt="Arrêter",
        key="micro",
        use_container_width=False,
    )
    st.caption("Cliquez puis parlez — arrêtez quand vous avez terminé")

if fichier:
    ext = fichier.name.lower().split(".")[-1]
    if ext == "txt":
        st.session_state["texte_dicte"] = fichier.read().decode("utf-8")
        st.session_state["last_image_name"] = None
    elif fichier.name != st.session_state.get("last_image_name"):
        st.session_state["last_image_name"] = fichier.name
        with st.spinner("Lecture du document en cours..."):
            try:
                mime = "image/png" if ext == "png" else "image/jpeg"
                st.session_state["texte_dicte"] = transcribe_image(fichier.getvalue(), mime)
                st.success(f"Texte extrait de {fichier.name}")
            except Exception as e:
                st.error(f"Erreur extraction : {e}")
else:
    if audio:
        with st.spinner("Transcription en cours..."):
            try:
                st.session_state["texte_dicte"] = transcribe_audio(audio["bytes"])
                st.success("Transcription terminée.")
            except Exception as e:
                st.error(f"Erreur transcription : {e}")

texte = st.text_area(
    "Texte",
    value=st.session_state.get("texte_dicte", ""),
    height=220,
    placeholder="Collez un extrait de texte ici\n— ou téléchargez un fichier .txt\n— ou dictez avec le micro\n— ou importez un scan JPG/PNG",
    label_visibility="collapsed",
)

st.divider()

# ── Section Paramètres ────────────────────────────────────────────────────────
st.markdown('<div class="aschool-section-title">Paramètres de l\'activité</div>', unsafe_allow_html=True)

matiere = st.session_state.get("matiere", "Français")
ACTIVITES = ACTIVITES_PAR_MATIERE.get(matiere, ACTIVITES_PAR_MATIERE["Français"])

col_act, col_niv = st.columns([1, 1])
with col_act:
    activite_label = st.selectbox("Type d'activité", list(ACTIVITES.keys()))
with col_niv:
    niveau = st.selectbox("Niveau de la classe", ["6e", "5e", "4e", "3e", "2nde", "1ère", "Terminale", "Supérieur"], index=2)

activite_cfg = ACTIVITES[activite_label]
activite_key = activite_cfg["key"]

sous_type = None
nb = None

col_st, col_nb = st.columns([1, 1])
with col_st:
    if activite_cfg["sous_types"]:
        sous_type = st.selectbox("Précision", activite_cfg["sous_types"])
with col_nb:
    if "nb" in activite_cfg["params"]:
        nb = st.number_input("Nombre de questions", min_value=1, value=5, step=1)

avec_correction = st.checkbox(
    "Inclure une proposition de correction",
    value=False,
    help="A-SCHOOL génère une réponse-type après chaque question, que le professeur adapte à sa classe.",
)

st.divider()

generer = st.button(
    "Générer l'activité",
    type="primary",
    use_container_width=True,
    help="Lancer la génération de l'activité avec A-SCHOOL",
)

if generer:
    if not texte or not texte.strip():
        st.error("Veuillez coller un texte, importer un fichier ou utiliser la dictée.")
    else:
        with st.spinner("Génération en cours..."):
            try:
                kwargs = {"niveau": niveau}
                if sous_type:
                    kwargs["sous_type"] = sous_type
                if nb:
                    kwargs["nb"] = nb

                prompt = build_prompt(activite_key, texte, avec_correction=avec_correction, **kwargs)
                resultat = generate(prompt)
                md_content = to_markdown(activite_key, niveau, resultat)
                today = date.today().strftime("%Y%m%d")

                Path("outputs").mkdir(exist_ok=True)
                save(md_content, f"outputs/{activite_key}_{today}.md")

                st.session_state["resultat"] = resultat
                st.session_state["resultat_key"] = activite_key
                st.session_state["resultat_today"] = today

            except Exception as e:
                st.error(f"Erreur : {e}")

# ── Section Résultat ──────────────────────────────────────────────────────────
if st.session_state.get("resultat"):
    resultat = st.session_state["resultat"]
    activite_key_dl = st.session_state["resultat_key"]
    today_dl = st.session_state["resultat_today"]

    st.divider()
    st.markdown('<div class="aschool-section-title">Résultat généré</div>', unsafe_allow_html=True)

    col_dl1, col_dl2, col_fermer = st.columns([2, 2, 1])
    with col_dl1:
        st.download_button(
            label="Télécharger en Word (.docx)",
            data=to_docx(resultat),
            file_name=f"{activite_key_dl}_{today_dl}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            help="Télécharger le résultat au format Word",
        )
    with col_dl2:
        st.download_button(
            label="Télécharger en texte (.txt)",
            data=resultat.encode("utf-8"),
            file_name=f"{activite_key_dl}_{today_dl}.txt",
            mime="text/plain",
            use_container_width=True,
            help="Télécharger le résultat en fichier texte",
        )
    with col_fermer:
        if st.button("Fermer", use_container_width=True, help="Fermer le résultat"):
            st.session_state["resultat"] = None
            st.rerun()

    st.markdown(f'<div class="aschool-result">{resultat}</div>', unsafe_allow_html=True)
